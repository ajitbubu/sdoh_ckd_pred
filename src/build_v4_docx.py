"""
src/build_v4_docx.py
--------------------
Convert CKD_Paper_JMIR_PHS_v4_revised.md to a JMIR-formatted .docx with all
13 figures embedded inline at first mention and core tables embedded in the
Results section.

JMIR Public Health & Surveillance formatting conventions applied:
  • Times New Roman 12 pt body, 1.5 line spacing
  • Title 16 pt bold, centered
  • Section headings 14 pt bold, left-aligned, blue accent
  • Sub-headings 12 pt bold
  • Figure captions BELOW figures, italic, 10 pt
  • Table titles ABOVE tables, bold, 11 pt
  • 1-inch margins
  • Single-column layout (JMIR submission format)
"""

import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH  = os.path.join(BASE_DIR, "CKD_Paper_JMIR_PHS_v4_revised.md")
TBL_DIR  = os.path.join(BASE_DIR, "tables_phs")
FIG_DIR  = os.path.join(BASE_DIR, "figures_phs")
OUT_PATH = os.path.join(BASE_DIR, "CKD_Paper_JMIR_PHS_v4_revised.docx")


# ── Figure registry: maps figure number → (filename, caption) ────────────
FIGURES = {
    1: ("fig01_parity.png",
        "Figure 1. Tract-level parity plot of model-predicted vs CDC PLACES-observed CKD prevalence (n=60,445 census tracts). Each point is one census tract; the dashed line is perfect agreement. Cross-validated R² = 0.450."),
    2: ("fig02_stage1_calibration.png",
        "Figure 2. Calibration plot of the Stage 1 ecological model across all 60,445 tracts grouped into 10 deciles of predicted prevalence. Calibration slope and intercept are reported in the figure."),
    3: ("fig03_loro_panels.png",
        "Figure 3. Leave-one-Census-region-out cross-validation. Each panel shows R², MAE, and calibration slope when the named region is held out and the model is trained on the other three regions. Heterogeneity: Northeast (R²=0.25) and West (R²=0.24) are harder to predict than Midwest or South."),
    4: ("fig04_capacity_sweep.png",
        "Figure 4. Sensitivity of cross-validated R² to XGBoost capacity. Five settings spanning n_estimators=200–1000, max_depth=3–7. R² varies by less than 0.005, indicating the model is not over-fit to a single hyperparameter choice."),
    5: ("fig05_adi_quintile.png",
        "Figure 5. Predicted (model) versus observed (CDC PLACES) CKD prevalence stratified by Area Deprivation Index quintile (1 = least deprived, 5 = most deprived). Both series show a monotonic gradient consistent with published deprivation–CKD effects."),
    6: ("fig06_feature_importance.png",
        "Figure 6. XGBoost gain importance for the three ADI features used in the Stage 1 ecological model (national rank, state rank, quintile)."),
    7: ("fig07_performance_matrix.png",
        "Figure 7. Performance matrix for the Stage 1 ecological model — overall and per-region cross-validated R², MAE, and calibration slope."),
    8: ("fig08_roc.png",
        "Figure 8. Receiver Operating Characteristic curve for the Stage 2 NHANES patient-level CKD classifier (out-of-fold predictions, 5-fold stratified CV with NHANES MEC weights). Survey-weighted AUROC = 0.773 (95% bootstrap CI 0.756–0.789)."),
    9: ("fig09_pr.png",
        "Figure 9. Precision–Recall curve for the Stage 2 NHANES classifier. Survey-weighted AUPRC = 0.404; horizontal line shows the survey-weighted CKD prevalence baseline (0.139)."),
    10: ("fig10_calibration.png",
         "Figure 10. Calibration plot for the Stage 2 NHANES classifier across deciles of predicted probability. Calibration slope = 0.955, intercept = +0.015."),
    11: ("fig11_subgroup_forest.png",
         "Figure 11. Forest plot of subgroup AUROC for the Stage 2 NHANES classifier across age band, sex, and race/ethnicity categories, with 95% bootstrap confidence intervals (1,000 replicates). Vertical dashed line = overall AUROC."),
    12: ("fig12_decision_curve.png",
         "Figure 12. Decision-curve analysis (Vickers 2006) for the Stage 2 NHANES classifier. Net benefit of the model exceeds the treat-all and treat-none strategies across the threshold range of 0.05–0.40."),
    13: ("fig13_performance_matrix.png",
         "Figure 13. Performance matrix for the Stage 2 NHANES classifier — overall metrics, calibration, subgroup AUROC range, and operating-point sensitivity."),
}

# ── Table registry: figure number → (csv path, title) ────────────────────
TABLES = {
    "Table 1": ("table1_baseline_characteristics.csv",
                "Table 1. Baseline characteristics of the NHANES analytic sample (n=15,150), survey-weighted, by CKD status."),
    "Table 2": ("table2_performance_summary.csv",
                "Table 2. Performance metrics summary for the Stage 2 NHANES classifier."),
    "Table 3a": ("table3a_confusion_matrix.csv",
                 "Table 3a. Confusion matrix at the 90%-specificity operating point (survey-weighted, in millions)."),
    "Table 3b": ("table3b_operating_metrics.csv",
                 "Table 3b. Operating-point metrics at the 90%-specificity threshold."),
}


# ── JMIR styling helpers ─────────────────────────────────────────────────
JMIR_BODY_FONT = "Times New Roman"
JMIR_BODY_SIZE = Pt(12)
JMIR_LINE_SP   = 1.5

NAVY  = RGBColor(0x1F, 0x4E, 0x79)
GREY  = RGBColor(0x55, 0x55, 0x55)


def jmir_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = JMIR_BODY_FONT
    style.font.size = JMIR_BODY_SIZE
    style.paragraph_format.line_spacing = JMIR_LINE_SP
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    return doc


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.bold = True
    run.font.name = JMIR_BODY_FONT
    run.font.size = Pt(16)


def add_meta(doc, line):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(line)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY


def add_h(doc, text, level):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = NAVY if level <= 2 else None


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = JMIR_LINE_SP
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.0)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.4)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.4)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def add_figure(doc, fig_num):
    """Embed a figure with caption below, JMIR style."""
    if fig_num not in FIGURES:
        return
    filename, caption = FIGURES[fig_num]
    fig_path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(fig_path):
        add_body(doc, f"_[Figure {fig_num} placeholder — file missing: {filename}]_")
        return
    # Add the image, centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(6.0))
    # Caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = 1.15
    crun = cap.add_run(caption)
    crun.italic = True
    crun.font.size = Pt(10)
    crun.font.color.rgb = GREY


def add_table_csv(doc, table_key):
    """Embed a CSV table with title above, JMIR style."""
    csv_name, title = TABLES[table_key]
    csv_path = os.path.join(TBL_DIR, csv_name)
    if not os.path.exists(csv_path):
        add_body(doc, f"_[{table_key} placeholder — file missing: {csv_name}]_")
        return
    # Title above
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(12)
    tp.paragraph_format.space_after = Pt(4)
    trun = tp.add_run(title)
    trun.bold = True
    trun.font.size = Pt(11)
    # Table
    df = pd.read_csv(csv_path)
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            cell = table.cell(i + 1, j)
            cell.text = "" if pd.isna(v) else str(v)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table


# ── Section-boundary insertion: when we encounter a section header, ─────
# insert any figures/tables registered for the section we are LEAVING.
# This guarantees figures 1-7 appear in Stage 1 results, figures 8-13 +
# tables 1-3 in Stage 2 results, in numerical order.
SECTION_INSERT_BEFORE = {
    # Before Stage 2 results section, dump Stage 1 figures
    "3.2 Stage 2 — Patient-level NHANES classifier": [
        ("fig", 1), ("fig", 2), ("fig", 3), ("fig", 4),
        ("fig", 5), ("fig", 6), ("fig", 7),
    ],
    # Before Discussion, dump Stage 2 figures + tables
    "4. Discussion": [
        ("table", "Table 1"),
        ("fig", 8), ("fig", 9), ("fig", 10),
        ("fig", 11), ("fig", 12), ("fig", 13),
        ("table", "Table 2"),
        ("table", "Table 3a"),
        ("table", "Table 3b"),
    ],
}


def emit_section_assets(doc, section_text):
    """When we hit a header that matches, insert its registered assets BEFORE."""
    for trigger, items in SECTION_INSERT_BEFORE.items():
        if section_text.strip() == trigger:
            for kind, key in items:
                if kind == "fig":
                    add_figure(doc, key)
                elif kind == "table":
                    add_table_csv(doc, key)
            return


# ── Main converter ───────────────────────────────────────────────────────
def main():
    with open(MD_PATH) as f:
        md = f.read()

    doc = jmir_doc()
    lines = md.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line == "---":
            i += 1
            continue

        # Skip raw figure-caption lines from the MD because we render them via add_figure
        if re.match(r"^\*\*Figure \d+\.\*\*", line):
            i += 1
            continue
        if re.match(r"^\*\*Table \d", line):
            i += 1
            continue

        # Title (single # in our doc)
        if line.startswith("# "):
            add_title(doc, line[2:].strip())
            i += 1
            continue

        # H2
        if line.startswith("## "):
            heading = line[3:].strip()
            emit_section_assets(doc, heading)
            add_h(doc, heading, 1)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            heading = line[4:].strip()
            emit_section_assets(doc, heading)
            add_h(doc, heading, 2)
            i += 1
            continue

        # H4
        if line.startswith("#### "):
            add_h(doc, line[5:].strip(), 3)
            i += 1
            continue

        # Meta lines under title
        if line.startswith("**Submission target:**") or \
           line.startswith("**Manuscript version:**") or \
           line.startswith("**Change summary:**"):
            add_meta(doc, line.replace("**", ""))
            i += 1
            continue

        # Bullets
        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        # Numbered
        m = re.match(r"^(\d+)\. (.+)$", line)
        if m:
            add_numbered(doc, m.group(2))
            i += 1
            continue

        # Skip raw markdown tables (we embed CSVs instead)
        if line.startswith("| "):
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                i += 1
            continue

        if not line:
            i += 1
            continue

        # Default: body paragraph
        add_body(doc, line)
        i += 1

    doc.save(OUT_PATH)
    print(f"Saved → {OUT_PATH}")
    print(f"  Size: {os.path.getsize(OUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
